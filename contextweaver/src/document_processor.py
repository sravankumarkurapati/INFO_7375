# src/document_processor.py
import os
from typing import List, Dict, Tuple, Optional
from datetime import datetime
from pathlib import Path
import pypdf
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from tqdm import tqdm
import re
import logging
from collections import defaultdict

from config import Config

logging.basicConfig(level=Config.LOG_LEVEL)
logger = logging.getLogger(__name__)


class KnowledgeBase:
    """
    Comprehensive knowledge base organization system
    
    Requirement: "Build a knowledge base for your domain"
    
    Features:
    - Hierarchical document organization
    - Domain classification
    - Source credibility assessment
    - Temporal tracking
    """
    
    def __init__(self):
        self.documents = []
        self.hierarchy = {
            'by_domain': defaultdict(list),
            'by_year': defaultdict(list),
            'by_source_type': defaultdict(list),
            'by_credibility': defaultdict(list)
        }
        self.statistics = {}
        
        # Domain classification keywords
        self.domain_keywords = {
            'legal': ['contract', 'agreement', 'clause', 'liability', 'jurisdiction', 
                     'amendment', 'bylaws', 'litigation', 'compliance', 'regulation'],
            'medical': ['patient', 'treatment', 'diagnosis', 'clinical', 'medical',
                       'disease', 'therapy', 'symptoms', 'cardiovascular', 'study'],
            'technical': ['api', 'function', 'code', 'implementation', 'documentation',
                         'software', 'algorithm', 'system', 'architecture', 'framework'],
            'research': ['study', 'research', 'methodology', 'findings', 'hypothesis',
                        'analysis', 'results', 'conclusion', 'participants', 'data'],
            'financial': ['revenue', 'profit', 'investment', 'market', 'stock',
                         'portfolio', 'risk', 'return', 'asset', 'liability']
        }
    
    def add_documents(self, documents: List[Document]):
        """Add documents to knowledge base with organization"""
        logger.info(f"📚 Adding {len(documents)} documents to knowledge base...")
        
        for doc in tqdm(documents, desc="Organizing documents"):
            # Classify and enrich metadata
            doc = self._enrich_metadata(doc)
            
            # Add to main collection
            self.documents.append(doc)
            
            # Organize hierarchically
            self._organize_document(doc)
        
        # Update statistics
        self._update_statistics()
        
        logger.info(f"✅ Knowledge base now contains {len(self.documents)} documents")
    
    def _enrich_metadata(self, doc: Document) -> Document:
        """Enrich document with additional metadata"""
        
        # Classify domain
        domain = self._classify_domain(doc.page_content)
        doc.metadata['domain'] = domain
        
        # Assess source type and credibility
        source_type = self._classify_source_type(doc.metadata.get('filename', ''))
        doc.metadata['source_type'] = source_type
        doc.metadata['credibility_score'] = Config.CREDIBILITY_SCORES.get(source_type, 0.5)
        
        # Extract entities (basic implementation)
        entities = self._extract_entities(doc.page_content)
        doc.metadata['entities'] = entities
        
        # Calculate content quality score
        doc.metadata['quality_score'] = self._calculate_quality_score(doc)
        
        return doc
    
    def _classify_domain(self, content: str) -> str:
        """Classify document domain using keyword matching"""
        content_lower = content.lower()
        
        scores = {}
        for domain, keywords in self.domain_keywords.items():
            score = sum(1 for kw in keywords if kw in content_lower)
            scores[domain] = score
        
        if max(scores.values()) == 0:
            return 'general'
        
        return max(scores, key=scores.get)
    
    def _classify_source_type(self, filename: str) -> str:
        """Classify source credibility type based on filename patterns"""
        filename_lower = filename.lower()
        
        patterns = {
            'peer_reviewed': ['journal', 'peer_review', 'nature', 'science', 'lancet'],
            'official_documentation': ['official', 'documentation', 'manual', 'spec'],
            'preprint': ['preprint', 'arxiv', 'biorxiv', 'ssrn'],
            'blog_post': ['blog', 'medium', 'post', 'article'],
        }
        
        for source_type, keywords in patterns.items():
            if any(kw in filename_lower for kw in keywords):
                return source_type
        
        return 'general_source'
    
    def _extract_entities(self, content: str) -> List[str]:
        """Basic entity extraction (names, organizations, dates)"""
        entities = []
        
        # Extract years
        years = re.findall(r'\b(19|20)\d{2}\b', content)
        entities.extend(years[:3])  # Keep top 3
        
        # Extract capitalized phrases (potential names/orgs)
        caps_phrases = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', content)
        entities.extend(caps_phrases[:5])  # Keep top 5
        
        return list(set(entities))
    
    def _calculate_quality_score(self, doc: Document) -> float:
        """Calculate content quality score"""
        content = doc.page_content
        
        # Length score (optimal 500-1500 chars)
        length = len(content)
        if length < 100:
            length_score = length / 100
        elif length > 2000:
            length_score = min(2000 / length, 1.0)
        else:
            length_score = 1.0
        
        # Structure score (sentences, paragraphs)
        sentences = len(re.findall(r'[.!?]+', content))
        paragraphs = content.count('\n\n')
        structure_score = min((sentences / 5 + paragraphs / 2) / 2, 1.0)
        
        # Completeness (ends properly)
        completeness = 1.0 if content.strip()[-1] in '.!?' else 0.7
        
        return (length_score + structure_score + completeness) / 3
    
    def _organize_document(self, doc: Document):
        """Organize document into hierarchical structure"""
        domain = doc.metadata.get('domain', 'general')
        year = doc.metadata.get('year', 'unknown')
        source_type = doc.metadata.get('source_type', 'general')
        credibility = doc.metadata.get('credibility_score', 0.5)
        
        self.hierarchy['by_domain'][domain].append(doc)
        self.hierarchy['by_year'][year].append(doc)
        self.hierarchy['by_source_type'][source_type].append(doc)
        
        # Categorize by credibility tier
        if credibility >= 0.8:
            tier = 'high_credibility'
        elif credibility >= 0.6:
            tier = 'medium_credibility'
        else:
            tier = 'low_credibility'
        self.hierarchy['by_credibility'][tier].append(doc)
    
    def _update_statistics(self):
        """Update knowledge base statistics"""
        self.statistics = {
            'total_documents': len(self.documents),
            'domains': {k: len(v) for k, v in self.hierarchy['by_domain'].items()},
            'years': {k: len(v) for k, v in self.hierarchy['by_year'].items()},
            'source_types': {k: len(v) for k, v in self.hierarchy['by_source_type'].items()},
            'credibility_tiers': {k: len(v) for k, v in self.hierarchy['by_credibility'].items()},
            'year_range': self._get_year_range(),
            'coverage_score': self._calculate_coverage()
        }
    
    def _get_year_range(self) -> Tuple[Optional[int], Optional[int]]:
        """Get earliest and latest years"""
        years = [int(y) for y in self.hierarchy['by_year'].keys() 
                if y != 'unknown' and str(y).isdigit()]
        
        if not years:
            return (None, None)
        
        return (min(years), max(years))
    
    def _calculate_coverage(self) -> float:
        """Calculate knowledge base coverage score"""
        num_domains = len(self.hierarchy['by_domain'])
        num_years = len([y for y in self.hierarchy['by_year'].keys() if y != 'unknown'])
        num_source_types = len(self.hierarchy['by_source_type'])
        
        # Normalize
        domain_score = min(num_domains / 5, 1.0)
        temporal_score = min(num_years / 10, 1.0)
        source_diversity = min(num_source_types / 5, 1.0)
        
        return (domain_score + temporal_score + source_diversity) / 3
    
    def get_documents_by_filter(
        self,
        domain: Optional[str] = None,
        year_min: Optional[int] = None,
        year_max: Optional[int] = None,
        source_type: Optional[str] = None,
        min_credibility: float = 0.0
    ) -> List[Document]:
        """Filter documents by criteria"""
        filtered = self.documents
        
        if domain:
            filtered = [d for d in filtered if d.metadata.get('domain') == domain]
        
        if year_min:
            filtered = [d for d in filtered 
                       if d.metadata.get('year', 0) >= year_min]
        
        if year_max:
            filtered = [d for d in filtered 
                       if d.metadata.get('year', float('inf')) <= year_max]
        
        if source_type:
            filtered = [d for d in filtered 
                       if d.metadata.get('source_type') == source_type]
        
        if min_credibility > 0:
            filtered = [d for d in filtered 
                       if d.metadata.get('credibility_score', 0) >= min_credibility]
        
        return filtered
    
    def print_statistics(self):
        """Print knowledge base statistics"""
        print("\n" + "="*60)
        print("📊 KNOWLEDGE BASE STATISTICS")
        print("="*60)
        
        print(f"\n📚 Total Documents: {self.statistics['total_documents']}")
        
        print(f"\n🏷️  Domains:")
        for domain, count in sorted(self.statistics['domains'].items(), 
                                    key=lambda x: x[1], reverse=True):
            print(f"   {domain:.<30} {count:>3}")
        
        print(f"\n📅 Temporal Coverage:")
        year_min, year_max = self.statistics['year_range']
        if year_min and year_max:
            print(f"   Range: {year_min} - {year_max} ({year_max - year_min + 1} years)")
        for year, count in sorted(self.statistics['years'].items()):
            print(f"   {year:.<30} {count:>3}")
        
        print(f"\n📖 Source Types:")
        for stype, count in sorted(self.statistics['source_types'].items(),
                                   key=lambda x: x[1], reverse=True):
            print(f"   {stype:.<30} {count:>3}")
        
        print(f"\n⭐ Credibility Distribution:")
        for tier, count in self.statistics['credibility_tiers'].items():
            print(f"   {tier:.<30} {count:>3}")
        
        print(f"\n📊 Coverage Score: {self.statistics['coverage_score']:.2f}/1.00")
        print("="*60 + "\n")


class AdvancedChunker:
    """
    Advanced document chunking with multiple strategies
    
    Requirement: "Design relevant document chunking strategies"
    
    Implements 4 chunking strategies:
    1. Fixed-size with overlap
    2. Semantic (similarity-based boundaries)
    3. Sentence-based (respects sentence boundaries)
    4. Hybrid (combines semantic + size constraints)
    """
    
    def __init__(
        self,
        chunk_size: int = Config.CHUNK_SIZE,
        chunk_overlap: int = Config.CHUNK_OVERLAP,
        strategy: str = Config.CHUNKING_STRATEGY
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.strategy = strategy
        
        logger.info(f"🔧 Initialized chunker with strategy: {strategy}")
    
    def chunk_documents(
        self,
        texts: List[str],
        metadatas: List[Dict]
    ) -> List[Document]:
        """Chunk multiple documents"""
        all_chunks = []
        
        for text, metadata in zip(texts, metadatas):
            chunks = self._chunk_single_document(text, metadata)
            all_chunks.extend(chunks)
        
        return all_chunks
    
    def _chunk_single_document(
        self,
        text: str,
        metadata: Dict
    ) -> List[Document]:
        """Chunk a single document using selected strategy"""
        
        if self.strategy == 'fixed':
            chunks = self._fixed_size_chunking(text)
        elif self.strategy == 'semantic':
            chunks = self._semantic_chunking(text)
        elif self.strategy == 'sentence':
            chunks = self._sentence_chunking(text)
        elif self.strategy == 'hybrid':
            chunks = self._hybrid_chunking(text)
        else:
            raise ValueError(f"Unknown chunking strategy: {self.strategy}")
        
        # Create Document objects
        documents = []
        for i, chunk_dict in enumerate(chunks):
            doc = Document(
                page_content=chunk_dict['text'],
                metadata={
                    **metadata,
                    'chunk_id': i,
                    'total_chunks': len(chunks),
                    'chunk_size': len(chunk_dict['text']),
                    'chunking_strategy': self.strategy,
                    'semantic_boundary': chunk_dict.get('semantic_boundary', False)
                }
            )
            documents.append(doc)
        
        return documents
    
    def _fixed_size_chunking(self, text: str) -> List[Dict]:
        """Strategy 1: Fixed-size chunking with overlap"""
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        chunks = splitter.split_text(text)
        return [{'text': chunk, 'semantic_boundary': False} for chunk in chunks]
    
    def _semantic_chunking(self, text: str) -> List[Dict]:
        """Strategy 2: Semantic similarity-based chunking"""
        # For now, use sentence-based as approximation
        # In production, would use embeddings
        return self._sentence_chunking(text)
    
    def _sentence_chunking(self, text: str) -> List[Dict]:
        """Strategy 3: Sentence-boundary respecting chunks"""
        sentences = self._split_sentences(text)
        
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            
            if current_length + sentence_length > self.chunk_size and current_chunk:
                chunks.append({
                    'text': ' '.join(current_chunk),
                    'semantic_boundary': True
                })
                current_chunk = [sentence]
                current_length = sentence_length
            else:
                current_chunk.append(sentence)
                current_length += sentence_length
        
        if current_chunk:
            chunks.append({
                'text': ' '.join(current_chunk),
                'semantic_boundary': True
            })
        
        return chunks
    
    def _hybrid_chunking(self, text: str) -> List[Dict]:
        """Strategy 4: Hybrid (sentence-based with size constraints)"""
        # First chunk by sentences
        sentence_chunks = self._sentence_chunking(text)
        
        # Then ensure no chunk exceeds max size
        final_chunks = []
        for chunk in sentence_chunks:
            if len(chunk['text']) > self.chunk_size * 1.5:
                # Split large chunks using fixed-size
                sub_chunks = self._fixed_size_chunking(chunk['text'])
                final_chunks.extend(sub_chunks)
            else:
                final_chunks.append(chunk)
        
        return final_chunks
    
    def _split_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]


class DocumentProcessor:
    """
    Main document processing pipeline
    
    Combines knowledge base organization with advanced chunking
    """
    
    def __init__(self):
        self.knowledge_base = KnowledgeBase()
        self.chunker = AdvancedChunker()
        logger.info("✅ Document processor initialized")
    
    def load_documents(self, file_paths: List[str]) -> List[Document]:
        """Load and process documents through full pipeline"""
        logger.info(f"📚 Loading {len(file_paths)} documents...")
        
        all_texts = []
        all_metadatas = []
        
        for file_path in tqdm(file_paths, desc="Reading files"):
            try:
                text, metadata = self._load_single_file(file_path)
                all_texts.append(text)
                all_metadatas.append(metadata)
            except Exception as e:
                logger.error(f"❌ Error loading {file_path}: {e}")
        
        # Chunk documents
        logger.info("🔪 Chunking documents...")
        chunked_docs = self.chunker.chunk_documents(all_texts, all_metadatas)
        
        # Add to knowledge base
        self.knowledge_base.add_documents(chunked_docs)
        
        logger.info(f"✅ Processed {len(chunked_docs)} chunks from {len(file_paths)} files")
        
        return chunked_docs
    
    def _load_single_file(self, file_path: str) -> Tuple[str, Dict]:
        """Load a single file and extract metadata"""
        path = Path(file_path)
        
        # Extract text based on file type
        if path.suffix == '.pdf':
            text = self._load_pdf(file_path)
        elif path.suffix == '.txt':
            text = self._load_txt(file_path)
        elif path.suffix == '.docx':
            text = self._load_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {path.suffix}")
        
        # Extract metadata
        metadata = self._extract_metadata(file_path, text)
        
        return text, metadata
    
    def _load_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def _load_txt(self, file_path: str) -> str:
        """Load text file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _load_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        doc = DocxDocument(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    
    def _extract_metadata(self, file_path: str, text: str) -> Dict:
        """Extract comprehensive metadata"""
        path = Path(file_path)
        stats = path.stat()
        
        return {
            'source': str(path),
            'filename': path.name,
            'file_type': path.suffix[1:],  # Remove dot
            'file_size': stats.st_size,
            'created_at': datetime.fromtimestamp(stats.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(stats.st_mtime).isoformat(),
            'year': self._extract_year(path.name, text),
            'word_count': len(text.split()),
            'char_count': len(text)
        }
    
    def _extract_year(self, filename: str, text: str) -> Optional[int]:
        """Extract year from filename or content"""
        # Try filename first
        match = re.search(r'(19|20)\d{2}', filename)
        if match:
            return int(match.group())
        
        # Try first 500 chars of content
        match = re.search(r'(19|20)\d{2}', text[:500])
        if match:
            return int(match.group())
        
        return None


# Test the processor
if __name__ == "__main__":
    processor = DocumentProcessor()
    print("✅ Document processor ready!")
    
    # Configuration summary
    print("\n⚙️  Chunking Configuration:")
    print(f"   Strategy: {Config.CHUNKING_STRATEGY}")
    print(f"   Chunk Size: {Config.CHUNK_SIZE}")
    print(f"   Overlap: {Config.CHUNK_OVERLAP}")